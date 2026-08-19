"""Render the per-page transcripts into one txt / docx / pdf buffer."""

import io
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

from ..core.font_registry import UNICODE_FONT_REGISTERED


def _page_sort_key(x: str):
    """Sort key that handles numeric ('5') and split page keys ('5a', '5b')."""
    match = re.match(r'^(\d+)(.*)', x)
    if match:
        return (int(match.group(1)), match.group(2))
    return (0, x)


def _page_display_label(key: str) -> str:
    return f"Page {key}"


def build_combined_transcript(transcripts: dict) -> str:
    """One string, pages in order, separated by a rule."""
    pages = sorted(transcripts.keys(), key=_page_sort_key)

    sections = []
    separator = '\u2500' * 20
    for page in pages:
        text = transcripts[page]
        label = _page_display_label(page)
        section = f"{label}\n{separator}\n{text}"
        sections.append(section)

    return "\n\n".join(sections)


def build_txt_export(transcripts: dict) -> io.BytesIO:
    """UTF-8 with a BOM, so Windows Notepad/Excel don't mangle accents."""
    combined = build_combined_transcript(transcripts)
    utf8_bom = b'\xef\xbb\xbf'
    buffer = io.BytesIO(utf8_bom + combined.encode('utf-8'))
    buffer.seek(0)
    return buffer


def build_docx_export(transcripts: dict) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading("Combined Transcript", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    pages = sorted(transcripts.keys(), key=_page_sort_key)

    for i, page in enumerate(pages):
        text = transcripts[page]

        doc.add_heading(_page_display_label(page), level=1)

        separator = doc.add_paragraph("─" * 40)
        separator.runs[0].font.size = Pt(10)

        for para_text in text.split('\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                para.style.font.size = Pt(11)

        if i < len(pages) - 1:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf_export(transcripts: dict) -> io.BytesIO:
    """Needs DejaVu Sans registered, or accents fall back to Helvetica."""
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    styles = getSampleStyleSheet()

    font_name = 'DejaVuSans' if UNICODE_FONT_REGISTERED else 'Helvetica'

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=18,
        spaceAfter=30,
        alignment=1  # center
    )

    page_header_style = ParagraphStyle(
        'PageHeader',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        spaceAfter=6,
        textColor='#1e40af'
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=11,
        spaceAfter=6,
        leading=14
    )

    separator_style = ParagraphStyle(
        'Separator',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=10,
        spaceAfter=12,
        textColor='#6b7280'
    )

    story = []
    story.append(Paragraph("Combined Transcript", title_style))
    story.append(Spacer(1, 20))

    pages = sorted(transcripts.keys(), key=_page_sort_key)

    for page in pages:
        text = transcripts[page]

        story.append(Paragraph(_page_display_label(page), page_header_style))
        story.append(Paragraph("─" * 50, separator_style))

        # reportlab parses a mini-HTML in Paragraph, so escape the text first.
        for line in text.split('\n'):
            if line.strip():
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, body_style))
            else:
                story.append(Spacer(1, 6))

        story.append(Spacer(1, 30))

    doc.build(story)
    buffer.seek(0)
    return buffer
